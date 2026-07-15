"""Build AlphaMaster beginner Word tutorial with embedded screenshots."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
IMG = ROOT / "docs" / "tutorial_images"
OUT = ROOT / "docs" / "AlphaMaster使用教程.docx"


def set_run_font(run, *, size=11, bold=False, color=None, name="微软雅黑"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    if color is not None:
        run.font.color.rgb = color
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), name)


def add_title(doc: Document, text: str):
    p = doc.add_heading(text, level=0)
    for run in p.runs:
        set_run_font(run, size=22, bold=True)


def add_h1(doc: Document, text: str):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        set_run_font(run, size=16, bold=True)


def add_h2(doc: Document, text: str):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        set_run_font(run, size=13, bold=True)


def add_p(doc: Document, text: str, *, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35
    return p


def add_steps(doc: Document, items: list[str]):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_run_font(run, size=11)
        p.paragraph_format.space_after = Pt(4)


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=11)
        p.paragraph_format.space_after = Pt(3)


def add_tip(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run("小贴士：")
    set_run_font(run, size=11, bold=True, color=RGBColor(0x0B, 0x6E, 0x4F))
    run2 = p.add_run(text)
    set_run_font(run2, size=11, color=RGBColor(0x1F, 0x3A, 0x33))
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)


def add_warn(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run("注意：")
    set_run_font(run, size=11, bold=True, color=RGBColor(0xB4, 0x3B, 0x2E))
    run2 = p.add_run(text)
    set_run_font(run2, size=11)
    p.paragraph_format.space_after = Pt(8)


def add_img(doc: Document, name: str, *, width_cm: float = 15.5, caption: str | None = None):
    path = IMG / name
    if not path.exists():
        add_p(doc, f"（缺少截图：{name}）", bold=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption)
        set_run_font(r, size=9, color=RGBColor(0x55, 0x65, 0x70))
        c.paragraph_format.space_after = Pt(10)


def build() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    add_title(doc, "AlphaMaster 超详细使用教程（带图）")
    add_p(doc, "适合人群：第一次接触本软件的同学（尽量写得像说明书）。", bold=True)
    add_p(
        doc,
        "这份教程会带你走完完整流程：安装与启动 → 准备数据 → 训练挖因子 → 回测检验 → 实时看信号。",
    )
    add_p(doc, f"对应软件路径示例：D:\\cl\\AlphaMaster　　网页地址：http://127.0.0.1:8765")
    add_p(doc, "文档生成日期：以你电脑上打开的软件界面为准（本教程插图来自实际截图）。")

    # 目录式大纲
    add_h1(doc, "目录")
    add_bullets(
        doc,
        [
            "一、这个软件是干什么的？",
            "二、开始前要准备什么",
            "三、怎么安装依赖（第一次必做）",
            "四、怎么启动软件",
            "五、打开后页面长什么样",
            "六、准备 K 线数据文件（最关键！）",
            "七、步骤① 训练：让电脑帮你挖因子",
            "八、步骤② 回测：用历史数据检验策略",
            "九、步骤③ 实时分析：盯盘出信号",
            "十、常用按钮一览表",
            "十一、遇到报错怎么办",
            "十二、建议的最佳实践",
        ],
    )

    add_h1(doc, "一、这个软件是干什么的？")
    add_p(
        doc,
        "一句话：AlphaMaster 会用人工智能（强化学习）在历史行情里“挖”交易因子公式，再帮你回测和实时观察信号。",
    )
    add_p(doc, "你可以把它想成下面三步：")
    add_steps(
        doc,
        [
            "训练：喂给它一份 K 线数据，它拼命搜索更好的公式。",
            "回测：把找到的公式放到历史行情上验成绩（看赚不赚钱、稳不稳）。",
            "实时分析：盯着当前市场，按同一套公式报方向（涨/跌/观望）。",
        ],
    )
    add_tip(doc, "你不需要会写代码。多数时间只需点按钮、选文件、看图表。")

    add_h1(doc, "二、开始前要准备什么")
    add_bullets(
        doc,
        [
            "一台 Windows 电脑（本教程按 Windows 写）。",
            "已安装 Python 3.10 及以上（推荐 3.11 或 3.13）。",
            "AlphaMaster 项目文件夹（例如 D:\\cl\\AlphaMaster 或 D:\\Al_master）。",
            "至少一份行情 Parquet 文件（后面教你命名）。",
            "如果要用 MT5 实时数据：电脑上还要打开并登录 MetaTrader 5。",
            "如果要用 AI 分析：准备 DeepSeek Key，或本机已开 QClaw / OpenClaw 一类本地助手。",
        ],
    )

    add_h1(doc, "三、怎么安装依赖（第一次必做）")
    add_p(doc, "打开「命令提示符」或 PowerShell，进入项目文件夹，输入：")
    add_p(doc, "cd /d D:\\cl\\AlphaMaster", bold=True)
    add_p(doc, "python -m pip install -r requirements.txt", bold=True)
    add_warn(
        doc,
        "请用 python -m pip，不要直接敲 pip（中文 Windows 上容易因编码报错）。",
    )
    add_p(doc, "装完可用下面一行检查关键包是否在：")
    add_p(
        doc,
        'python -c "import torch,fastapi,MetaTrader5,pyarrow,matplotlib,multipart; print(\'OK\')"',
        bold=True,
    )
    add_tip(
        doc,
        "TradingView 数据源若安装失败，可再执行："
        "python -m pip install git+https://github.com/rongardF/tvdatafeed.git",
    )

    add_h1(doc, "四、怎么启动软件")
    add_h2(doc, "方法 A（推荐）：双击启动脚本")
    add_steps(
        doc,
        [
            "进入项目文件夹，找到 start_web.bat。",
            "双击它。脚本会：先清理被占用的端口 → 启动服务 → 等页面真正就绪 → 自动打开浏览器。",
            "看到浏览器打开 http://127.0.0.1:8765 就成功了。",
            "启动器窗口可以先留着；关掉启动器窗口一般不会立刻关服务（服务在另一个“AlphaMaster Server”窗口）。",
        ],
    )
    add_h2(doc, "方法 B：命令行启动")
    add_p(doc, "cd /d D:\\cl\\AlphaMaster", bold=True)
    add_p(doc, "python run_web.py --host 127.0.0.1 --port 8765", bold=True)
    add_p(doc, "然后在浏览器地址栏输入：http://127.0.0.1:8765")
    add_warn(
        doc,
        "如果提示端口被占用，可以改成 --port 8766，或先关闭旧的 AlphaMaster Server 窗口。",
    )

    add_h1(doc, "五、打开后页面长什么样")
    add_p(doc, "打开网页后，最上方是软件标题，中间有三个大步骤：01 模型训练、02 策略回测、03 实时分析。")
    add_img(doc, "01_home_top.png", caption="图1 打开软件后的首页上半部分")
    add_img(doc, "02_stepper.png", width_cm=14.5, caption="图2 顶部三步导航：训练 → 回测 → 实时")
    add_p(doc, "你要使用时，就从上到下按这三个步骤走。一般不要跳着乱点。")

    add_h1(doc, "六、准备 K 线数据文件（最关键！）")
    add_p(
        doc,
        "本软件训练和回测强制使用本地 Parquet 文件，不会在训练时偷偷联网拉行情。所以「文件名正确」非常重要。",
    )
    add_h2(doc, "6.1 文件怎么命名")
    add_p(doc, "标准格式：", bold=True)
    add_p(doc, "{品种}_{周期}.parquet", bold=True)
    add_p(doc, "正确例子：")
    add_bullets(
        doc,
        [
            "XAUUSD_H1.parquet（黄金 · 1小时）",
            "BTCUSDT_H1.parquet（比特币 · 1小时）",
            "AAPL_H1.parquet（苹果股票 · 1小时）",
            "002008_60min.parquet（A股代码 · 60分钟，软件会自动认成 H1）",
            "600519_5min.parquet（会认成 M5）",
        ],
    )
    add_p(doc, "周期别名对照（你会用到的）：", bold=True)
    add_bullets(
        doc,
        [
            "H1 = 1小时 = 60min / 60m / 1h",
            "M5 = 5分钟 = 5min / 5m",
            "M15 = 15分钟 = 15min / 15m",
            "H4 = 4小时 = 4h / 240min",
            "D1 = 日线 = 1d / day / daily",
        ],
    )
    add_h2(doc, "6.2 文件里面要有哪些列")
    add_p(doc, "Parquet 至少要有：time、open、high、low、close，以及 volume 或 tick_volume。")
    add_p(doc, "K 线数量也不能太少，否则会提示「数据不足」。")
    add_warn(doc, "不要选 Excel、CSV 直接训练。必须是 .parquet。")

    add_h1(doc, "七、步骤① 训练：让电脑帮你挖因子")
    add_p(doc, "确保顶部导航停在「01 模型训练」。")

    add_h2(doc, "7.1 选择数据文件")
    add_img(doc, "03_train_launch.png", caption="图3 训练启动区：选择数据文件 + 各类按钮")
    add_steps(
        doc,
        [
            "点击「选择数据文件」。",
            "在弹出的文件窗口里，找到你的 xxx_H1.parquet（或 002008_60min.parquet）。",
            "选中后，中间卡片会显示品种、周期、K线数量。如果写「文件不存在」或红色报错，先检查路径和命名。",
            "只有选对文件后，「开始训练」才会变成可点。",
        ],
    )

    add_h2(doc, "7.2 开始训练")
    add_steps(
        doc,
        [
            "点击蓝色「开始训练」。右上角状态会从「空闲」变成训练中。",
            "左侧会出现「训练曲线」：绿色是最优分数，蓝色是验证分数。",
            "右侧「训练日志」会不断刷出进度。",
            "下方「最优公式」会显示目前找到的最好因子写法。",
        ],
    )
    add_img(doc, "04_train_chart.png", caption="图4 训练曲线区域（开始训练后会动）")
    add_img(doc, "05_train_log.png", caption="图5 训练日志区域")
    add_img(doc, "06_train_full.png", width_cm=14.0, caption="图6 训练页整页示意")

    add_h2(doc, "7.3 两个分数分别是什么意思（很重要）")
    add_bullets(
        doc,
        [
            "最优分数：到目前为止挖到的「最好那一条公式」的成绩。通常只升不降，抬一截说明挖到更好的了。",
            "验证分数：最近一批公式在验证集上的平均表现，会上下波动，属正常。",
            "不要只因为验证分数某一刻掉了就慌——软件有时会主动「重新搅一搅」去探索新公式。",
        ],
    )

    add_h2(doc, "7.4 「开始训练」和「重新训练」有什么区别")
    add_bullets(
        doc,
        [
            "开始训练：接着上次的检查点继续挖（省时间）。",
            "重新训练：清空检查点，从零重新搜一轮（可能挖到更好的，但更花时间）。",
            "已有更好策略不会被随便冲掉：只有新公式更强才会覆盖。",
        ],
    )

    add_h2(doc, "7.5 什么时候可以停")
    add_bullets(
        doc,
        [
            "最优分数长时间几乎不动，可以考虑停止。",
            "点红色「停止」。停止后，最好策略一般已保存在 strategies 文件夹。",
            "也可以继续挂着挖，时间越长通常机会越多（但费电费时间）。",
        ],
    )

    add_h2(doc, "7.6 导出策略 / 导出训练 / 导入训练")
    add_bullets(
        doc,
        [
            "导出策略：下载当前品种最优策略 JSON，方便备份或给别人回测。",
            "导出训练：把检查点、曲线、策略打成 zip，换电脑可继续。",
            "导入训练：上传以前导出的 zip 或 .pt，下次可断点续训。",
        ],
    )

    add_h2(doc, "7.7 已保存策略列表")
    add_img(doc, "07_strategies.png", caption="图7 已保存策略表：能看到品种、周期、分数、公式")
    add_p(doc, "训练成功产生的策略会出现在这里。回测和实时分析经常会用到这些文件。")

    add_h2(doc, "7.8 AI 分析（可选）")
    add_img(doc, "08_ai_panel.png", caption="图8 AI 分析面板")
    add_steps(
        doc,
        [
            "若使用 DeepSeek：在输入框填入 API Key。",
            "若本机有 QClaw / OpenClaw：可按界面提示切换 provider。",
            "点击「开始分析」，AI 会用白话解释：现在训练情况如何、值不值得继续、因子大概在干什么。",
        ],
    )
    add_tip(doc, "AI 分析是助手，不是买卖建议。最终仍要用回测结果自己判断。")

    add_h1(doc, "八、步骤② 回测：用历史数据检验策略")
    add_p(doc, "点顶部「02 策略回测」。")
    add_img(doc, "09_backtest_top.png", caption="图9 回测页顶部")
    add_img(doc, "10_backtest_launch.png", caption="图10 选择策略、设置手续费/滑点、开始回测")

    add_h2(doc, "8.1 怎么做回测")
    add_steps(
        doc,
        [
            "点击「选择策略」，选 strategies 里的 best_品种.json，或你导出的策略 JSON。",
            "（可选）改手续费、滑点。默认单边手续费 0.02%、滑点 0.01%，合计约 0.03%。",
            "确认该策略能找到对应 Parquet（策略里常记录 data_file；没有则回退训练页同品种文件）。",
            "点「开始回测」，等日志跑完。",
            "查看「回测绩效」「绩效明细」和下方资金曲线。",
        ],
    )
    add_img(doc, "11_backtest_summary.png", caption="图11 回测绩效区域（跑完后会出数字）")
    add_img(doc, "12_backtest_full.png", width_cm=14.0, caption="图12 回测页整页示意")

    add_h2(doc, "8.2 回测结果怎么读（小学生版）")
    add_bullets(
        doc,
        [
            "收益：这段历史上是赚还是亏。",
            "Sharpe / Sortino：稳不稳；越高一般越好（但也要结合交易次数看）。",
            "盈亏比：赚的时候平均赚多少，对比亏的时候平均亏多少。",
            "胜率：猜对方向的比例。胜率高不一定就好，还要看盈亏比。",
            "交易数：太少不够信；多了也不等于一定稳。",
        ],
    )
    add_warn(
        doc,
        "回测好看 ≠ 未来一定赚钱。至少换一段数据、或换手续费成本再验一次更踏实。",
    )
    add_warn(
        doc,
        "若报错 No module named 'matplotlib'：执行 python -m pip install matplotlib 后重试。",
    )

    add_h1(doc, "九、步骤③ 实时分析：盯盘出信号")
    add_p(doc, "点顶部「03 实时分析」。")
    add_img(doc, "13_realtime_top.png", caption="图13 实时分析页顶部")
    add_img(doc, "14_realtime_form.png", caption="图14 添加监控：数据源 + 品种 + 周期 + 策略因子")
    add_img(doc, "15_realtime_full.png", width_cm=14.0, caption="图15 实时分析整页示意")

    add_h2(doc, "9.1 添加一个监控怎么操作")
    add_steps(
        doc,
        [
            "选择数据源：常见是 MT5 或 TradingView（界面当前主要展示这两者）。",
            "填写品种，例如 XAUUSD、EURUSD。也可从下拉预设里选。",
            "选择周期，尽量和策略训练时的周期一致（例如策略是 H1，这里也选 1h）。",
            "选择策略因子：选本机 best_xxx.json，或点导入策略。",
            "点击添加。页面会出现一张「信号卡片」。",
        ],
    )

    add_h2(doc, "9.2 信号卡片怎么看")
    add_bullets(
        doc,
        [
            "预期上涨 / 预期下跌 / 先观望：当前公式给出的方向。",
            "把握大小：强度翻译成白话，告诉你这拨信号有多坚定。",
            "距离下次判断：还有多久重新算一次。",
            "若状态是「错误」：看卡片上的红字，通常是数据源连不上或历史不够。",
        ],
    )

    add_h2(doc, "9.3 TradingView 连不上怎么办")
    add_p(
        doc,
        "如果添加 TradingView 监控时弹出「无法使用 TradingView」，按弹窗提示做：",
    )
    add_bullets(
        doc,
        [
            "把 VPN 设成全局，并打开 TUN（虚拟网卡）模式再试。",
            "或按弹窗「使用云服务器」去看部署说明。",
            "或改回 MT5 数据源（本机 MT5 已登录时更稳）。",
        ],
    )

    add_h2(doc, "9.4 飞书提醒（可选）")
    add_p(
        doc,
        "实时页可配置飞书机器人：当方向发生转折时推送文本提醒。按页面上的帮助去开通 webhook 即可。",
    )

    add_h1(doc, "十、常用按钮一览表")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "按钮"
    hdr[1].text = "在哪一页"
    hdr[2].text = "干什么"
    rows = [
        ("选择数据文件", "训练", "挑选本地 Parquet"),
        ("开始训练", "训练", "继续挖因子"),
        ("重新训练", "训练", "清空检查点重搜"),
        ("停止", "训练/回测", "停掉当前任务"),
        ("导出策略", "训练", "下载最优策略 JSON"),
        ("导出/导入训练", "训练", "迁移训练进度"),
        ("开始分析", "训练·AI区", "让 AI 解释训练情况"),
        ("选择策略", "回测", "挑选要检验的策略"),
        ("开始回测", "回测", "跑历史绩效"),
        ("添加监控", "实时", "挂一个实时信号任务"),
    ]
    for a, b, c in rows:
        cells = table.add_row().cells
        cells[0].text = a
        cells[1].text = b
        cells[2].text = c
    doc.add_paragraph()

    add_h1(doc, "十一、遇到报错怎么办")
    add_h2(doc, "11.1 页面一打开就一堆「网络错误 Failed to fetch」")
    add_bullets(
        doc,
        [
            "通常是服务刚启动，还在加载 torch，还没准备好。",
            "请用 start_web.bat 启动（它会等接口就绪再开浏览器）。",
            "新版本前端会自动重试几次；仍失败就等 10 秒刷新页面。",
        ],
    )
    add_h2(doc, "11.2 文件名报错（例如 002008_60min.parquet）")
    add_p(
        doc,
        "新版本已支持 60min/5min/1h 等别名。若仍报错，请确认你用的是已更新的 parquet_manager，并已重启 Web。",
    )
    add_h2(doc, "11.3 回测说没有 matplotlib")
    add_p(doc, "python -m pip install matplotlib", bold=True)
    add_h2(doc, "11.4 选择数据文件提示缺 pyarrow")
    add_p(doc, "python -m pip install pyarrow", bold=True)
    add_h2(doc, "11.5 历史训练时长一上来就很大")
    add_p(
        doc,
        "可能是仓库自带作者训练记录。清空 training_time_品种.json，并删除 logs/train_品种_*.log 后再刷新。",
    )
    add_h2(doc, "11.6 MT5 数据源不可用")
    add_bullets(
        doc,
        [
            "确认已安装 MetaTrader5 Python 包。",
            "确认 MT5 终端已打开并登录账号。",
        ],
    )

    add_h1(doc, "十二、建议的最佳实践（照着做不容易踩坑）")
    add_steps(
        doc,
        [
            "一次只训练一个品种、一个周期（别同时开好多任务抢 CPU）。",
            "文件名写清楚：品种_周期.parquet。",
            "训练时多看「最优分数」是否还在抬升。",
            "训完先回测，再考虑实时。",
            "实时监控的周期尽量与策略训练周期一致。",
            "定期导出策略和训练包做备份。",
            "换电脑部署：复制项目 → python -m pip install -r requirements.txt → start_web.bat。",
        ],
    )

    add_h1(doc, "附录：最短上手路线（10 分钟版）")
    add_steps(
        doc,
        [
            "双击 start_web.bat，等浏览器打开。",
            "准备好 XAUUSD_H1.parquet（或你自己的数据）。",
            "训练页 → 选择数据文件 → 开始训练。",
            "过一阵点停止（或等分数抬升后停）。",
            "回测页 → 选择刚生成的策略 → 开始回测。",
            "实时页 → 选 MT5 → 填同一品种与周期 → 选同一策略 → 添加。",
        ],
    )
    add_p(doc, "到这里，你已经走完 AlphaMaster 的完整闭环。", bold=True)
    add_p(
        doc,
        "插图目录：docs/tutorial_images/　　本文档：docs/AlphaMaster使用教程.docx",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    return OUT


if __name__ == "__main__":
    path = build()
    print("saved", path)
